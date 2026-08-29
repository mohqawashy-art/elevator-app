"""نموذج عقد Word: رفع من الإعدادات وتعبئة الحقول المتغيرة فقط."""
from __future__ import annotations

import os
from datetime import date
from io import BytesIO

from docx import Document

from app import app, db
from models import Contract, Customer, Settings
from tests.conftest import login_as


def _make_template_bytes() -> bytes:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run('طرف ثاني: ')
    p.add_run('{{اسم_العميل}}')
    p.add_run(' — عقد ')
    p.add_run('{{رقم_العقد}}')
    p.add_run(' بقيمة ')
    p.add_run('{{قيمة_العقد}}')
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_settings_shows_contract_template_upload(client):
    login_as(client, 'admin')
    html = client.get('/settings').get_data(as_text=True)
    assert 'نموذج عقد Word' in html
    assert 'اسم_العميل' in html
    assert 'contract_template' in html


def test_fill_contract_docx_keeps_placeholder_values(client, tmp_path):
    login_as(client, 'admin')
    from contract_docx import fill_contract_docx, placeholder_map

    with app.app_context():
        org_id = Settings.query.first().organization_id
        cust = Customer(
            code='C-WD01',
            name='مؤسسة الاختبار',
            national_id='1012345678',
            status='نشط',
            organization_id=org_id,
        )
        db.session.add(cust)
        db.session.flush()
        contract = Contract(
            code='CN-WD01',
            customer_id=cust.id,
            contract_type='صيانة',
            start_date=date(2026, 1, 1),
            end_date=date(2026, 12, 31),
            duration_months=12,
            value=24000,
            total=27600,
            status='نشط',
            organization_id=org_id,
        )
        db.session.add(contract)
        db.session.commit()
        mapping = placeholder_map(contract)
        assert mapping['{{اسم_العميل}}'] == 'مؤسسة الاختبار'
        assert mapping['{{رقم_الهوية}}'] == '1012345678'

        tpl = tmp_path / 'cn.docx'
        tpl.write_bytes(_make_template_bytes())
        filled = fill_contract_docx(contract, str(tpl))

    out = Document(BytesIO(filled))
    text = '\n'.join(p.text for p in out.paragraphs)
    assert 'مؤسسة الاختبار' in text
    assert '{{اسم_العميل}}' not in text
    assert '{{رقم_العقد}}' not in text


def test_contract_print_downloads_docx_when_template_set(client):
    login_as(client, 'admin')
    with app.app_context():
        org_id = Settings.query.first().organization_id
        cust = Customer(code='C-WD02', name='عميل ورد', status='نشط', organization_id=org_id)
        db.session.add(cust)
        db.session.flush()
        contract = Contract(
            code='CN-WD02',
            customer_id=cust.id,
            contract_type='صيانة',
            start_date=date.today(),
            end_date=date.today(),
            value=1000,
            total=1150,
            status='نشط',
            organization_id=org_id,
        )
        db.session.add(contract)
        s = Settings.query.first()
        dest_dir = os.path.join(app.root_path, 'static', 'uploads', 'company', 'test-docx')
        os.makedirs(dest_dir, exist_ok=True)
        path = os.path.join(dest_dir, 'contract-template-test.docx')
        with open(path, 'wb') as fh:
            fh.write(_make_template_bytes())
        s.contract_template_path = 'uploads/company/test-docx/contract-template-test.docx'
        db.session.commit()
        cid = contract.id

    r = client.get(f'/contracts/{cid}/print')
    assert r.status_code == 200
    disp = r.headers.get('Content-Disposition', '')
    ctype = r.headers.get('Content-Type', '')
    assert 'wordprocessingml' in ctype or '.docx' in disp
    assert r.get_data()[:2] == b'PK'

    html = client.get(f'/contracts/{cid}/print?html=1')
    assert html.status_code == 200
    assert 'طباعة' in html.get_data(as_text=True)
